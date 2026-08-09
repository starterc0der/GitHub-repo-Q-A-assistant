from __future__ import annotations

from src.ingest.sources import chunk_csv, load_csv, load_docx, load_pdf, load_text


def test_load_text_wraps_pasted_string_as_one_logical_file() -> None:
    files = load_text("hello world", "Pasted note")
    assert len(files) == 1
    assert files[0].name == "Pasted note"
    assert files[0].text == "hello world"


def test_load_text_skips_blank_paste() -> None:
    assert load_text("   ", "Pasted note") == []


def test_load_docx_reads_paragraphs_as_one_logical_file(tmp_path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    path = tmp_path / "notes.docx"
    doc.save(path)

    files = load_docx(path, "notes.docx")

    assert len(files) == 1
    assert files[0].name == "notes.docx"
    assert "First paragraph." in files[0].text
    assert "Second paragraph." in files[0].text


def test_load_pdf_skips_pages_with_no_extractable_text(tmp_path) -> None:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    path = tmp_path / "blank.pdf"
    with open(path, "wb") as f:
        writer.write(f)

    assert load_pdf(path, "blank.pdf") == []


def test_load_csv_reads_whole_file_as_one_logical_file(tmp_path) -> None:
    path = tmp_path / "phones.csv"
    path.write_text("Name,Price\niPhone 16,799\niPhone 17,899\n")

    files = load_csv(path, "phones.csv")

    assert len(files) == 1
    assert files[0].name == "phones.csv"
    assert "iPhone 16,799" in files[0].text


def test_load_csv_skips_blank_file(tmp_path) -> None:
    path = tmp_path / "empty.csv"
    path.write_text("   \n")

    assert load_csv(path, "empty.csv") == []


def _csv_text(n_rows: int) -> str:
    header = "Name,Price"
    rows = "\n".join(f"Row{i},{i}" for i in range(1, n_rows + 1))
    return f"{header}\n{rows}\n"


def test_chunk_csv_stamps_header_onto_every_chunk_not_just_the_first() -> None:
    """The exact bug this was built to avoid: without this, only the first chunk would
    carry the column names and every later chunk would be bare, unlabeled numbers."""
    text = _csv_text(120)  # 3 chunks at the default 50 rows/chunk

    chunks = chunk_csv(text, "phones.csv", "space1", "src1", rows_per_chunk=50)

    assert len(chunks) == 3
    assert all(c.context_header == "Name,Price" for c in chunks)


def test_chunk_csv_never_splits_a_row_and_keeps_real_line_numbers() -> None:
    text = _csv_text(10)

    chunks = chunk_csv(text, "phones.csv", "space1", "src1", rows_per_chunk=4)

    assert len(chunks) == 3
    assert chunks[0].code == "Row1,1\nRow2,2\nRow3,3\nRow4,4"
    assert (chunks[0].start_line, chunks[0].end_line) == (2, 5)  # header is line 1
    assert chunks[1].code == "Row5,5\nRow6,6\nRow7,7\nRow8,8"
    assert (chunks[1].start_line, chunks[1].end_line) == (6, 9)
    assert chunks[2].code == "Row9,9\nRow10,10"
    assert (chunks[2].start_line, chunks[2].end_line) == (10, 11)


def test_chunk_csv_on_header_only_file_returns_no_chunks() -> None:
    assert chunk_csv("Name,Price\n", "phones.csv", "space1", "src1") == []
